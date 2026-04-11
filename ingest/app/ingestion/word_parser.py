from docx import Document
from pathlib import Path
from app.ingestion.base import BaseParser
from app.models.document import ParsedDocument, FileFormat
from app.core.exceptions import ParseFailureError


class WordParser(BaseParser):
    @property
    def supported_format(self) -> FileFormat:
        return FileFormat.WORD

    @property
    def supported_extensions(self) -> set[str]:
        return {"docx"}

    def parse(self, file_path: Path) -> ParsedDocument:
        try:
            doc = Document(file_path)
            parags = [p.text for p in doc.paragraphs if p.text.strip()]
            headings = [
                {"text": p.text} for p in doc.paragraphs if "Heading" in p.style.name
            ]

            props = doc.core_properties
            return ParsedDocument(
                file_name=file_path.name,
                format=self.supported_format,
                page_count=None,
                text_content=parags,
                title=props.title or None,
                author=props.author or None,
                subject=props.subject or None,
                keywords=props.keywords or None,
                creator=props.author or None,
                producer=None,
                creation_date=str(props.created) if props.created else None,
                mod_date=str(props.modified) if props.modified else None,
                metadata={"revision": props.revision, "paragraph_count": len(parags)},
            )

        except Exception as e:
            raise ParseFailureError(file_path.name, str(e))
